import os
from os import path
from docx import Document
import re

def main():
    print("FIRST\t" + os.getcwd())
    #os.chdir("Z Binder")
    print("NEXT\t" + os.getcwd())

    myDir = os.listdir(".")

    # for file in myDir:
    #     if file.endswith(".docx"):
    #         print("file name: " + file)
    #         name = os.path.splitext(os.path.basename(file))[0]
    #         print("file without extension: " + name)
    #
    #         doc = Document(file)
    #         print("Paragraph objects: ")
    #
    #         os.chdir("../ONSONG")
    #         lineNumber = 0
    #
    #         newFile = open(name + ".txt", "w", encoding="utf-8")
    #
    #
    #         for paragraph in doc.paragraphs:
    #             print("PARAGRAPH: " + paragraph.text)
    #
    #             empty = 1
    #
    #             for run in paragraph.runs:
    #
    #                 print("RUN : " + run.text)
    #
    #                 if lineNumber is 0:
    #                     newstring = re.sub(r"^[0-9]*\.*", "", paragraph.text)
    #                     print("\t\tREGEXED: " + newstring)
    #                     newFile.write("*\t\t" + newstring)
    #                     newFile.write("\nNumber: " + name)
    #                     lineNumber = lineNumber  + 1
    #                     break
    #
    #                 if run.bold:
    #                     #print("RUN IS BOLD " + paragraph.text)
    #                     newFile.write("*\t\t" + paragraph.text + "\n")
    #                     empty = 0
    #                     lineNumber = lineNumber + 1
    #                     break
    #                     #print("\t\t" + run.text)
    #
    #
    #                 else:
    #                     newFile.write(paragraph.text + "\n")
    #                     empty = 0
    #                     lineNumber = lineNumber + 1
    #                     break
    #                     #print("NOT " + i.text)
    #
    #             if empty == 1:
    #
    #                 newFile.write("\n")
    #     os.chdir("../Z Binder")


    os.chdir("CHORDS FOR ONSONG")
    myDir = os.listdir(".")

    for file in myDir:
        if file.endswith(".docx"):
            print("file name: " + file)
            name = os.path.splitext(os.path.basename(file))[0]
            print("file without extension: " + name)

            doc = Document(file)
            print("Paragraph objects: ")

            os.chdir("../ONSONG")
            newFile = open(name + ".txt", "w", encoding="utf-8")

            for paragraph in doc.paragraphs:
                print("PARAGRAPH: " + paragraph.text)

                empty = 1

                for run in paragraph.runs:

                    print("RUN : " + run.text)
                    if run.bold:
                        print("RUN IS BOLD " + paragraph.text)
                        newstring = re.sub(r"^[0-9]*\.*", "", paragraph.text)
                        print("\t\tREGEXED: " + newstring)
                        newFile.write("*                " + newstring + "\n")
                        empty = 0
                        break
                        #print("\t\t" + run.text)


                    else:
                        newFile.write(paragraph.text + "\n")
                        empty = 0
                        break
                        #print("NOT " + i.text)

                if empty == 1:

                    newFile.write("\n")
        os.chdir("../CHORDS FOR ONSONG")


if __name__ == "__main__":
    main()
